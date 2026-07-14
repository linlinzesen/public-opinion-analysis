# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

style = doc.styles['Normal']
style.font.name = '宋体'; style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5


def set_cell(cell, text, bold=False, size=Pt(10), align='left'):
    cell.text = ''
    p = cell.paragraphs[0]; r = p.add_run(text)
    r.font.name = '宋体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    r.font.size = size; r.bold = bold
    if align == 'center': p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = '黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h


def para(text, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.75)
    r = p.add_run(text)
    r.font.name = '宋体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    r.font.size = Pt(11); r.bold = bold


def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    r = p.add_run(text)
    r.font.name = 'Consolas'; r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(30, 30, 30)


# ═══════════ 封面 ═══════════
doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('网络舆情事件智能分析系统'); r.font.name = '黑体'
r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体'); r.font.size = Pt(26); r.bold = True

t2 = doc.add_paragraph(); t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run('—— 平台启动与使用说明 ——'); r2.font.name = '黑体'
r2._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体'); r2.font.size = Pt(16)
doc.add_paragraph()

# ═══════════ 一、环境要求 ═══════════
heading('一、环境要求', 1)

t = doc.add_table(rows=5, cols=3, style='Table Grid')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['模块', '运行环境', '要求']):
    set_cell(t.rows[0].cells[i], h, bold=True, align='center')
for r, d in enumerate([
    ['前端', 'Node.js + npm', 'Node.js ≥ 18'],
    ['后端', 'Python 虚拟环境 (.venv)', 'Python ≥ 3.10'],
    ['爬虫管道', 'Python 虚拟环境 (复用后端)', 'Python ≥ 3.10'],
    ['LLM 服务', 'Python 虚拟环境 (复用后端)', 'Python ≥ 3.10'],
]):
    for c, v in enumerate(d): set_cell(t.rows[r+1].cells[c], v)
doc.add_paragraph()

para('项目使用虚拟环境管理 Python 依赖，无需系统中全局安装 Python 包。后端、爬虫、LLM 共用 backend/.venv。前端使用 npm 管理 Node.js 依赖。')

# ═══════════ 二、启动前的准备 ═══════════
heading('二、启动前的准备', 1)

heading('2.1 安装 Python 依赖（首次或依赖变更时执行）', 2)
para('打开终端（cmd 或 PowerShell），进入项目根目录：')
code('cd backend')
code('.venv\\Scripts\\activate')
code('pip install -r requirements.txt')
para('安装完成后虚拟环境中将包含：Flask, flask-cors, PyJWT。')
para('爬虫及 NLP 相关依赖（如首次使用需要安装）：')
code('pip install pandas requests scikit-learn jieba snownlp')
doc.add_paragraph()

heading('2.2 安装前端依赖', 2)
code('cd frontend')
code('npm install')
doc.add_paragraph()

heading('2.3 准备数据（二选一）', 2)
para('方式一：运行爬虫流水线（生成数据）', bold=True)
code('cd crawler')
code('..\\backend\\.venv\\Scripts\\activate')
code('python run_pipeline.py')
para('流水线执行约 1~3 分钟，完成后在 crawler\\data\\ 下生成：')
code('  analysis_result.json    ← 后端要导入的文件')
code('  raw_comments.csv        ← 原始评论')
code('  labeled_comments.csv    ← 情感标注后评论')
para('方式二：跳过爬虫，直接用模拟数据', bold=True)
para('后端内置了兜底模拟数据，首次启动时自动初始化数据库，含 5 个预置舆情事件——不运行爬虫也可以直接体验前端。')
para('导入数据到数据库：')
code('cd backend')
code('.venv\\Scripts\\activate')
code('python data_import.py')

# ═══════════ 三、启动步骤 ═══════════
heading('三、启动步骤', 1)

para('推荐启动顺序：后端 → （可选 LLM）→ 前端。所有命令均在项目根目录下执行。')

heading('3.1 启动后端 API（端口 5000）★ 必须先启动', 2)
code('cd backend')
code('.venv\\Scripts\\activate          # 激活虚拟环境')
code('python app.py')
para('控制台看到以下信息说明启动成功：')
code('[app] 数据库已初始化')
code('[app] LLM 服务已挂载 → /api/llm/*')
code('[app] 服务启动 → http://localhost:5000')
para('验证：浏览器访问 http://localhost:5000/api/health 应返回 {"service":"opinion-backend"}。')

heading('3.2 启动 LLM 智能服务（端口 5001，可选）', 2)
para('后端 app.py 启动时已自动挂载 LLM 模块，通常无需单独启动此服务。')
para('如果后端控制台显示「LLM 服务未加载」，说明 LLM 模块路径配置有误，此时可独立启动：')
code('cd LLM/llm_service')
code('..\\..\\backend\\.venv\\Scripts\\activate')
para('如需真实 AI 能力，先设置环境变量再启动：')
code('set LLM_API_KEY=你的API密钥')
code('set LLM_MODEL=qwen-plus')
code('python app.py')
para('未配置 Key 时自动降级为 Mock 模板回答，不影响功能演示。')

heading('3.3 启动前端（端口 5173）★ 最后启动', 2)
code('cd frontend')
code('npm run dev')
para('控制台看到以下信息说明启动成功：')
code('VITE v5.x.x  ready in xxx ms')
code('→ Local:   http://localhost:5173/')
para('用浏览器打开 http://localhost:5173 即可访问系统。')

# ═══════════ 四、网站访问 ═══════════
heading('四、网站访问地址', 1)

para('全部启动后，在浏览器中打开：')
para('', bold=False)

# 用醒目的样式写出 URL
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('http://localhost:5173')
r.font.name = 'Consolas'; r.font.size = Pt(22); r.bold = True
r.font.color.rgb = RGBColor(37, 99, 235)

doc.add_paragraph()
para('输入默认账号登录：用户名 admin，密码 123456', bold=True)

t3 = doc.add_table(rows=2, cols=3, style='Table Grid')
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['用户名', '密码', '角色']):
    set_cell(t3.rows[0].cells[i], h, bold=True, align='center')
set_cell(t3.rows[1].cells[0], 'admin'); set_cell(t3.rows[1].cells[1], '123456')
set_cell(t3.rows[1].cells[2], '系统管理员')
doc.add_paragraph()

para('各服务端口一览：')
t4 = doc.add_table(rows=4, cols=3, style='Table Grid')
t4.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['服务', '端口', '地址']):
    set_cell(t4.rows[0].cells[i], h, bold=True, align='center')
for r, d in enumerate([
    ['前端页面', '5173', 'http://localhost:5173'],
    ['后端 API', '5000', 'http://localhost:5000/api'],
    ['LLM 服务', '5001', 'http://localhost:5001/api/llm'],
]):
    for c, v in enumerate(d): set_cell(t4.rows[r+1].cells[c], v)

# ═══════════ 五、功能验收清单 ═══════════
heading('五、功能验收清单', 1)

para('按以下步骤逐一验证各项功能：')
para('① 打开 http://localhost:5173 → 看到登录页', bold=True)
para('② 用 admin / 123456 登录 → 进入首页仪表盘', bold=True)
para('③ 首页应显示：5 个统计卡片、热度走势图、重点风险事件列表、AI 助手', bold=True)
para('④ 点击「事件看板」→ 5 个事件卡片列表，支持搜索/筛选/排序', bold=True)
para('⑤ 点击任一事件 → 详情页应包含：')
code('  · 事件基本信息（风险等级、情感倾向、热度、平台、时间）')
code('  · 热度趋势图 · 情感饼图 · 平台柱状图 · 高频词云')
code('  · 智能问答（输入问题 → 问 AI）')
code('  · 趋势预测（含生命周期阶段：潜伏期/成长期/高潮期/衰退期）')
code('  · 传播路径追踪（初始爆料 → 大V/官媒介入时间线）')
code('  · 信息可信度检测')
code('  · AI 报告生成（Markdown 格式）')
para('⑥ 点击「个人中心」→ 管理关注平台和关键词', bold=True)

# ═══════════ 六、API 速查 ═══════════
heading('六、核心 API 速查', 1)

para('所有接口需在 Header 中携带：Authorization: Bearer <token>')

t5 = doc.add_table(rows=13, cols=3, style='Table Grid')
t5.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['方法', '路径', '说明']):
    set_cell(t5.rows[0].cells[i], h, bold=True, align='center')
apis = [
    ['POST', '/api/auth/login', '登录获取 Token'],
    ['GET', '/api/overview/stats', '首页概览统计'],
    ['GET', '/api/events', '事件列表（keyword/riskLevel/sortBy）'],
    ['GET', '/api/events/:id', '事件详情'],
    ['GET', '/api/events/:id/trend', '热度趋势'],
    ['GET', '/api/events/:id/sentiment', '情感分布'],
    ['GET', '/api/events/:id/platforms', '平台分布'],
    ['GET', '/api/events/:id/word-cloud', '词云数据'],
    ['GET', '/api/events/:id/propagation', '★ 传播路径追踪'],
    ['POST', '/api/events/:id/credibility', '★ 可信度检测'],
    ['POST', '/api/llm/report', 'AI 生成报告'],
    ['POST', '/api/user/follow-platforms', '添加关注平台'],
]
for r, d in enumerate(apis):
    for c, v in enumerate(d): set_cell(t5.rows[r+1].cells[c], v)

# ═══════════ 七、FAQ ═══════════
heading('七、常见问题', 1)

heading('Q1: pip install 报错 / 模块找不到', 2)
para('确保已激活虚拟环境（终端提示符前应有 (.venv) 字样）：')
code('cd backend && .venv\\Scripts\\activate')
para('然后重新 pip install -r requirements.txt。')

heading('Q2: 前端页面打开但数据为空', 2)
para('检查后端是否已启动（http://localhost:5000/api/health）。检查 crawler\\data\\analysis_result.json 是否存在——如果不存在，运行 python data_import.py 或 python run_pipeline.py。')

heading('Q3: LLM 问答/报告返回模板回答', 2)
para('这是正常降级行为——未配置 LLM_API_KEY 时自动用模板兜底，不影响演示。如需真实 AI，设置环境变量后重启。')

heading('Q4: 爬虫运行报错', 2)
para('爬虫有自动降级机制——真实 API 不可用时生成逼真模拟数据，流水线始终能完成。如需真实数据，确保网络能访问 B站/微博/头条 API。')

heading('Q5: 端口被占用', 2)
para('5000/5173/5001 端口被占用时，修改对应配置文件中的端口号，并同步修改前端 .env 中的 API 地址。')

# 保存
out = os.path.join(os.path.dirname(os.path.abspath(__file__)), '平台启动与使用说明.docx')
doc.save(out)
print('OK: ' + out)
